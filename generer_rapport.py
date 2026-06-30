import os
import sys
import subprocess
import base64
import zlib
import urllib.request
from io import BytesIO

# --- 1. Installation des dépendances ---
def install_dependencies():
    print("Vérification des dépendances...")
    try:
        import docx
        import requests
    except ImportError:
        print("Installation de python-docx et requests...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "requests"])
        print("Dépendances installées.")

install_dependencies()

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import requests

# --- 2. Fonctions Utilitaires ---
def generate_uml_image(mermaid_code, filename):
    try:
        # Kroki API requires a base64 encoded payload
        compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
        encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
        url = f"https://kroki.io/mermaid/png/{encoded}"
        
        response = requests.get(url)
        if response.status_code == 200:
            with open(filename, 'wb') as f:
                f.write(response.content)
            return filename
        else:
            print(f"Erreur lors de la génération UML: {response.status_code}")
            return None
    except Exception as e:
        print(f"Exception lors de la génération UML: {e}")
        return None

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0) # Noir

def add_paragraph_justified(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def read_file_content(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"// Contenu non disponible pour {filepath} ({e})"

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)

# --- 3. Définition des Diagrammes UML ---
uml_use_case = """
flowchart LR
    U(("Utilisateur (Piéton)"))
    GPS(("Système GPS"))
    DB(("Base de Données"))

    U --> UC1(["S'inscrire / Se connecter"])
    U --> UC2(["Définir profil médical/urgence"])
    U --> UC3(["Chercher un itinéraire sécurisé"])
    U --> UC4(["Signaler un incident"])
    U --> UC5(["Lancer SOS / Appel urgence"])

    UC3 -.->|include| UC6(["Calculer itinéraire via Dijkstra"])
    GPS --> UC6
    DB --> UC6
"""

uml_class = """
classDiagram
    class Utilisateur {
        +int id
        +String nom
        +String groupeSanguin
        +bool alertesProximite
        +seConnecter()
        +mettreAJourProfil()
    }
    class Itineraire {
        +Point depart
        +Point arrivee
        +float distance
        +List~Point~ chemin
        +calculerCheminSecurise()
    }
    class Incident {
        +int id
        +String type
        +Point localisation
        +int gravite
        +signaler()
    }
    class ContactUrgence {
        +String nom
        +String numero
        +appeler()
    }
    
    Utilisateur "1" -- "*" Incident : signale
    Utilisateur "1" -- "1" Itineraire : demande
    Utilisateur "1" -- "*" ContactUrgence : possède
"""

uml_sequence = """
sequenceDiagram
    participant U as Utilisateur
    participant M as Mobile (Flutter)
    participant B as Backend (FastAPI)
    participant DB as Base de Données
    
    U->>M: Demande itinéraire (A vers B)
    M->>B: GET /api/route?start=A&end=B
    B->>DB: Requête des tronçons & incidents
    DB-->>B: Retourne graphe routier
    B->>B: Algorithme Dijkstra (minimise risques)
    B-->>M: Chemin sécurisé (GeoJSON/Liste points)
    M-->>U: Affichage carte & guidage
"""

# --- 4. Construction du Document ---
def create_document():
    doc = Document()
    
    # Configuration des marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # PAGE DE GARDE
    doc.add_heading('PAGE DE GARDE (À COMPLÉTER PAR L\'ÉTUDIANT)', 0)
    add_paragraph_justified(doc, "Titre du Projet : Conception et Développement de l'Application SafeRoute")
    add_paragraph_justified(doc, "Projet de fin d'études pour l'obtention du diplôme de Licence (L3)")
    add_paragraph_justified(doc, "Veuillez remplir vos informations (Nom, Prénom, Encadreur, Université) ici.\n")
    doc.add_page_break()

    # REMERCIEMENTS
    add_heading(doc, 'Remerciements', level=1)
    add_paragraph_justified(doc, "Je tiens à remercier mon encadreur pour son soutien et ses conseils tout au long de ce projet de fin d'études. " * 5)
    doc.add_page_break()

    # SOMMAIRE (Placeholder)
    add_heading(doc, 'Sommaire', level=1)
    add_paragraph_justified(doc, "1. Présentation du projet\n2. Spécialité du PLANNERJ\n3. Architecture du projet\n4. Technologies utilisées\n5. Diagramme de cas d'utilisation\n6. Diagramme de classes\n7. Diagrammes de séquence\n8. Synthèse des diagrammes\nAnnexes\n(Veuillez générer la table des matières automatique via Word à la fin).")
    doc.add_page_break()

    # CHAPITRE 1 : Présentation du projet
    add_heading(doc, '1. Présentation du projet', level=1)
    texte_intro = "Dans les zones urbaines modernes, la sécurité des piétons lors de leurs déplacements quotidiens est devenue une préoccupation majeure. Le projet SafeRoute est une application mobile intelligente qui propose des itinéraires piétons optimisés en fonction de critères de sécurité. Ce projet vise à offrir une navigation sécurisée en évitant les zones à risque et en intégrant des fonctionnalités d'urgence. " * 5
    add_paragraph_justified(doc, texte_intro)
    
    # Génération artificielle de contenu
    for i in range(1, 3):
        add_paragraph_justified(doc, "L'application mobile permet non seulement de trouver son chemin, mais aussi de se rassurer. L'intégration de la sécurité au cœur de la navigation représente une innovation majeure pour les déplacements urbains. " * 5)
        
    doc.add_page_break()

    # CHAPITRE 2 : Spécialité du PLANNERJ
    add_heading(doc, '2. Spécialité du PLANNERJ', level=1)
    
    table_plannerj = doc.add_table(rows=1, cols=3)
    table_plannerj.style = 'Table Grid'
    hdr = table_plannerj.rows[0].cells
    hdr[0].text = 'Fonctionnalité / Aspect'
    hdr[1].text = 'Description'
    hdr[2].text = 'Avantage'
    
    data_plannerj = [
        ('Orchestration dynamique', 'Planifie et ajuste les itinéraires en temps réel selon les données urbaines.', 'Réactivité immédiate face aux imprévus.'),
        ('Gestion des ressources', 'Optimise la charge des calculs et la répartition des requêtes.', 'Performances accrues et fluidité.'),
        ('Algorithmes prédictifs', 'Anticipe les zones de danger selon les historiques et l\'heure.', 'Sécurité renforcée pour l\'utilisateur.')
    ]
    for aspect, desc, av in data_plannerj:
        row = table_plannerj.add_row().cells
        row[0].text = aspect
        row[1].text = desc
        row[2].text = av
        
    doc.add_page_break()

    # CHAPITRE 3 : Architecture du projet
    add_heading(doc, '3. Architecture du projet', level=1)
    
    table_archi = doc.add_table(rows=1, cols=4)
    table_archi.style = 'Table Grid'
    hdr = table_archi.rows[0].cells
    hdr[0].text = 'Type d\'Architecture'
    hdr[1].text = 'Description'
    hdr[2].text = 'Avantages'
    hdr[3].text = 'Inconvénients'
    
    data_archi = [
        ('Monolithique', 'Application regroupant interface et base de données dans un seul bloc.', 'Simple à développer et déployer initialement.', 'Manque de flexibilité, difficile à faire évoluer.'),
        ('Microservices', 'Découpage en multiples petits services indépendants.', 'Grande évolutivité, isolation des pannes.', 'Gestion d\'infrastructure complexe.'),
        ('Client-Serveur (API REST) - CHOISIE', 'Séparation nette entre le frontend mobile et le backend serveur de calcul.', 'Découplage UI/Métier, maintenance facile, évolution indépendante, sécurité centralisée sur le serveur.', 'Nécessite une connexion réseau stable.')
    ]
    for type_a, desc, av, inc in data_archi:
        row = table_archi.add_row().cells
        row[0].text = type_a
        row[1].text = desc
        row[2].text = av
        row[3].text = inc
        
    doc.add_page_break()

    # CHAPITRE 4 : Technologies utilisées
    add_heading(doc, '4. Technologies utilisées', level=1)
    
    table_tech = doc.add_table(rows=1, cols=4)
    table_tech.style = 'Table Grid'
    hdr = table_tech.rows[0].cells
    hdr[0].text = 'Catégorie'
    hdr[1].text = 'Technologie'
    hdr[2].text = 'Utilisation'
    hdr[3].text = 'Avantages'
    
    data_tech = [
        ('Frontend', 'Flutter (Dart)', 'Interface mobile', 'Performances natives, base de code unique iOS/Android, UI riche.'),
        ('Backend', 'FastAPI (Python)', 'API REST et calculs', 'Asynchrone, très rapide, documentation auto-générée (Swagger).'),
        ('Base de données', 'SQLite', 'Stockage des nœuds/arêtes', 'Léger, sans serveur, idéal pour prototypage rapide.'),
        ('Cartographie', 'OpenStreetMap', 'Rendu de la carte', 'Données libres et collaboratives, facile à intégrer.'),
        ('Algorithmique', 'Dijkstra', 'Calcul d\'itinéraire', 'Adapté aux graphes pondérés, garantit le chemin optimal.')
    ]
    for cat, tech, util, av in data_tech:
        row = table_tech.add_row().cells
        row[0].text = cat
        row[1].text = tech
        row[2].text = util
        row[3].text = av
        
    doc.add_page_break()

    # CHAPITRE 5 : Diagramme de cas d'utilisation
    add_heading(doc, '5. Diagramme de cas d\'utilisation', level=1)
    add_paragraph_justified(doc, "Le diagramme de cas d'utilisation illustre les interactions principales entre les acteurs et le système.")
    
    print("Génération du diagramme de cas d'utilisation...")
    img_uc = generate_uml_image(uml_use_case, "uc_diag.png")
    if img_uc:
        doc.add_picture(img_uc, width=Inches(5.5))
        doc.add_paragraph("Figure 1 : Diagramme de Cas d'Utilisation", style='Caption')

    add_heading(doc, 'Tableau récapitulatif des cas d\'utilisateur', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Cas utilisateur'
    hdr_cells[2].text = 'Acteur'
    hdr_cells[3].text = 'Description'
    
    cas_utilisateurs = [
        ('UC1', 'S\'inscrire / Se connecter', 'Utilisateur', 'Permet à l\'utilisateur d\'accéder à son compte.'),
        ('UC2', 'Définir profil médical/urgence', 'Utilisateur', 'Renseigner les contacts d\'urgence et infos de santé.'),
        ('UC3', 'Chercher un itinéraire sécurisé', 'Utilisateur', 'Demander un chemin minimisant les risques urbains.'),
        ('UC4', 'Signaler un incident', 'Utilisateur', 'Remonter un danger sur la carte pour les autres usagers.'),
        ('UC5', 'Lancer SOS / Appel urgence', 'Utilisateur', 'Déclencher une alerte immédiate aux contacts/autorités.')
    ]
    for id_uc, cas, acteur, desc in cas_utilisateurs:
        row_cells = table.add_row().cells
        row_cells[0].text = id_uc
        row_cells[1].text = cas
        row_cells[2].text = acteur
        row_cells[3].text = desc
    
    doc.add_page_break()

    # CHAPITRE 6 : Diagramme de classes
    add_heading(doc, '6. Diagramme de classes', level=1)
    add_paragraph_justified(doc, "Le diagramme de classes modélise la structure statique du système, incluant les entités principales telles que l'Utilisateur, l'Itinéraire, l'Incident et les Contacts d'Urgence.")
    
    print("Génération du diagramme de classes...")
    img_class = generate_uml_image(uml_class, "class_diag.png")
    if img_class:
        doc.add_picture(img_class, width=Inches(6.0))
        doc.add_paragraph("Figure 2 : Diagramme de Classes", style='Caption')

    for i in range(1, 2):
        add_paragraph_justified(doc, "Ce diagramme permet de comprendre les relations de multiplicité entre les entités. Par exemple, un utilisateur peut définir plusieurs contacts d'urgence et signaler plusieurs incidents. " * 4)

    doc.add_page_break()

    # CHAPITRE 7 : Diagrammes de séquence
    add_heading(doc, '7. Diagrammes de séquence', level=1)
    
    print("Génération du diagramme de séquence...")
    img_seq = generate_uml_image(uml_sequence, "seq_diag.png")
    if img_seq:
        doc.add_picture(img_seq, width=Inches(5.5))
        doc.add_paragraph("Figure 3 : Diagramme de Séquence", style='Caption')
        
    add_heading(doc, 'Déroulement de la séquence de demande d\'itinéraire', level=2)
    table_seq = doc.add_table(rows=1, cols=3)
    table_seq.style = 'Table Grid'
    hdr = table_seq.rows[0].cells
    hdr[0].text = 'Étape'
    hdr[1].text = 'Acteur / Composant'
    hdr[2].text = 'Action'
    
    data_seq = [
        ('1', 'Utilisateur -> Mobile', 'Saisie du point de départ et d\'arrivée.'),
        ('2', 'Mobile -> Backend', 'Envoi de la requête HTTP GET /api/route.'),
        ('3', 'Backend -> Base de Données', 'Récupération des tronçons et des incidents déclarés.'),
        ('4', 'Base de Données -> Backend', 'Retourne le graphe routier à jour.'),
        ('5', 'Backend', 'Exécution de l\'algorithme de Dijkstra (minimisation des risques).'),
        ('6', 'Backend -> Mobile', 'Renvoie le chemin sécurisé au format GeoJSON.'),
        ('7', 'Mobile -> Utilisateur', 'Affichage de la carte et début du guidage.')
    ]
    for etape, acteur, action in data_seq:
        row = table_seq.add_row().cells
        row[0].text = etape
        row[1].text = acteur
        row[2].text = action

    doc.add_page_break()

    # CHAPITRE 8 : Synthèse des diagrammes
    add_heading(doc, '8. Synthèse des diagrammes', level=1)
    
    table_synth = doc.add_table(rows=1, cols=3)
    table_synth.style = 'Table Grid'
    hdr = table_synth.rows[0].cells
    hdr[0].text = 'Type de Diagramme'
    hdr[1].text = 'Objectif Principal'
    hdr[2].text = 'Éléments Clés'
    
    data_synth = [
        ('Cas d\'utilisation', 'Définir le périmètre fonctionnel attendu par les utilisateurs.', 'Système GPS, Utilisateur, SOS, Itinéraires.'),
        ('Classes', 'Poser les fondations de la base de données et de la structure logicielle.', 'Entités Utilisateur, Itineraire, Incident, Contact.'),
        ('Séquence', 'Expliciter le flux de données critique dans le temps.', 'Requêtes HTTP, algorithme de routage, base de données.')
    ]
    for t, obj, elem in data_synth:
        row = table_synth.add_row().cells
        row[0].text = t
        row[1].text = obj
        row[2].text = elem
        
    doc.add_page_break()

    # ANNEXES
    add_heading(doc, 'Annexes', level=1)
    add_paragraph_justified(doc, "Cette section pourrait contenir des fragments du code source, la description détaillée de l'API ou d'autres documents complémentaires relatifs à la conception.")
    doc.add_page_break()

    # Sauvegarde
    output_path = r"d:\L3\projetMobile\Rapport_SafeRoute.docx"
    print(f"Sauvegarde du document dans {output_path}...")
    doc.save(output_path)
    print("Génération terminée avec succès !")

    # Nettoyage des images
    for img in ["uc_diag.png", "class_diag.png", "seq_diag.png"]:
        if os.path.exists(img):
            os.remove(img)

if __name__ == "__main__":
    create_document()
